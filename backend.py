import re
import textwrap
from io import BytesIO
import fitz
import torch
import numpy as np
import nltk

from typing import Dict, Any
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, AutoModelForQuestionAnswering
from sklearn.metrics.pairwise import cosine_similarity

nltk.download("punkt")

EMBED_MODEL_NAME = "sentence-transformers/all-mpnet-base-v2"
SUMM_MODEL_NAME = "facebook/bart-large-cnn"
QA_MODEL_NAME = "distilbert-base-cased-distilled-squad"

device = "cuda" if torch.cuda.is_available() else "cpu"

metadata_store: Dict[str, Dict[str, Any]] = {}
embeddings_store = []

SUMMARY_LENGTH_CONFIG = {
    "short": {"max_length": 90, "min_length": 30},
    "medium": {"max_length": 200, "min_length": 160},
    "long": {"max_length": 360, "min_length": 200},
}


def extract_key_points(summary, max_points=3):

    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+", summary)
        if sentence.strip()
    ]

    if len(sentences) <= max_points:
        return sentences

    scored_sentences = [
        (idx, sentence, len(sentence.split()))
        for idx, sentence in enumerate(sentences)
    ]

    top_scored = sorted(
        scored_sentences,
        key=lambda item: item[2],
        reverse=True
    )[:max_points]

    # Keep original summary order for readability in UI.
    top_scored = sorted(top_scored, key=lambda item: item[0])

    return [sentence for _, sentence, _ in top_scored]


def compose_export_text(summary, key_points):

    lines = ["Generated Summary", "", summary, "", "Key Points"]

    for idx, point in enumerate(key_points, start=1):
        lines.append(f"{idx}. {point}")

    return "\n".join(lines)


def build_download_payload(summary, key_points, output_format):

    export_text = compose_export_text(summary, key_points)

    if output_format == "txt":
        return export_text.encode("utf-8"), "text/plain", "txt"

    if output_format == "word":
        from docx import Document

        doc = Document()
        doc.add_heading("Generated Summary", level=1)
        doc.add_paragraph(summary)
        doc.add_heading("Key Points", level=1)

        for point in key_points:
            doc.add_paragraph(point, style="List Number")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return (
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx"
        )

    if output_format == "pdf":
        page_width, page_height = fitz.paper_size("a4")
        margin = 50
        body_font_size = 11
        heading_font_size = 16
        section_font_size = 13
        line_step = 16

        pdf_doc = fitz.open()
        page = pdf_doc.new_page(width=page_width, height=page_height)
        y_pos = margin

        def add_line(text_line, font_size=body_font_size):
            nonlocal page, y_pos

            if y_pos > (page_height - margin):
                page = pdf_doc.new_page(width=page_width, height=page_height)
                y_pos = margin

            page.insert_text(
                (margin, y_pos),
                text_line,
                fontsize=font_size,
                fontname="helv"
            )
            y_pos += line_step

        add_line("Generated Summary", font_size=heading_font_size)
        y_pos += 8

        add_line("Summary", font_size=section_font_size)

        for summary_line in textwrap.wrap(summary, width=95):
            add_line(summary_line)

        y_pos += 8
        add_line("Key Points", font_size=section_font_size)

        for idx, point in enumerate(key_points, start=1):
            wrapped_point = textwrap.wrap(
                point,
                width=90,
                initial_indent=f"{idx}. ",
                subsequent_indent="   "
            )
            for point_line in wrapped_point:
                add_line(point_line)

        return pdf_doc.write(), "application/pdf", "pdf"

    raise ValueError("Unsupported format selected for download.")


# -----------------------------
# LOAD MODELS
# -----------------------------

def load_models():

    print("Loading models on", device)

    embed_model = SentenceTransformer(
        EMBED_MODEL_NAME,
        device=device
    )

    tokenizer = AutoTokenizer.from_pretrained(SUMM_MODEL_NAME)

    model = AutoModelForSeq2SeqLM.from_pretrained(
        SUMM_MODEL_NAME
    ).to(device)

    qa_tokenizer = AutoTokenizer.from_pretrained(QA_MODEL_NAME)
    qa_model = AutoModelForQuestionAnswering.from_pretrained(
        QA_MODEL_NAME
    ).to(device)

    return embed_model, tokenizer, model, (qa_tokenizer, qa_model)


# -----------------------------
# CLEAN TEXT
# -----------------------------

def clean_text(text):

    text = re.sub(r"\bPage\s*\d+\b", "", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s.,;:!?-]", "", text)

    return text.strip()


# -----------------------------
# CHUNK TEXT
# -----------------------------

def chunk_text(text, max_tokens=900, overlap=150):

    chunks = []
    start = 0

    while start < len(text):

        end = start + max_tokens
        chunks.append(text[start:end])
        start += (max_tokens - overlap)

    return chunks


# -----------------------------
# INGEST PDF
# -----------------------------

def ingest_pdf(file_path, paper_id, embed_model):

    global embeddings_store

    doc = fitz.open(file_path)

    text = " ".join([page.get_text() for page in doc])

    text = clean_text(text)

    chunks = chunk_text(text)

    embeddings = embed_model.encode(chunks)

    for i, emb in enumerate(embeddings):

        metadata_store[str(i)] = {
            "paper_id": paper_id,
            "text": chunks[i]
        }

    embeddings_store = embeddings

    print(f"{paper_id} indexed with {len(chunks)} chunks")


# -----------------------------
# RETRIEVE + SUMMARIZE
# -----------------------------

def smart_summarize(query, embed_model, tokenizer, model, summary_length="medium"):

    if len(embeddings_store) == 0:
        raise ValueError("No paper has been ingested for summarization.")

    generation_config = SUMMARY_LENGTH_CONFIG.get(
        summary_length,
        SUMMARY_LENGTH_CONFIG["medium"]
    )

    query_vec = embed_model.encode([query])

    sims = cosine_similarity(query_vec, embeddings_store)[0]

    top_k_idx = np.argsort(sims)[-3:]

    context = " ".join(
        metadata_store[str(i)]["text"] for i in top_k_idx
    )

    inputs = tokenizer(
        context,
        max_length=1024,
        truncation=True,
        return_tensors="pt"
    ).to(device)

    summary_ids = model.generate(
        inputs["input_ids"],
        max_length=generation_config["max_length"],
        min_length=generation_config["min_length"],
        num_beams=4,
        length_penalty=2.0,
        early_stopping=True
    )

    summary = tokenizer.decode(
        summary_ids[0],
        skip_special_tokens=True
    )

    key_points = extract_key_points(summary)

    return {
        "summary": summary,
        "key_points": key_points
    }


def _qa_single_chunk(question, context, qa_tokenizer, qa_model, max_answer_len=50):
    """Run extractive QA on a single chunk. Returns (answer, score, confidence)."""

    inputs = qa_tokenizer(
        question,
        context,
        return_offsets_mapping=True,
        return_tensors="pt",
        truncation="only_second",
        max_length=512,
    )

    offset_mapping = inputs.pop("offset_mapping")[0]
    sequence_ids = inputs.sequence_ids(0)

    model_inputs = {k: v.to(device) for k, v in inputs.items()}

    with torch.no_grad():
        outputs = qa_model(**model_inputs)

    start_logits = outputs.start_logits[0].cpu()
    end_logits = outputs.end_logits[0].cpu()

    # --- Find the best valid span ---
    best_span = None
    best_score = float("-inf")

    top_start = torch.topk(start_logits, k=min(20, start_logits.shape[0])).indices.tolist()
    top_end = torch.topk(end_logits, k=min(20, end_logits.shape[0])).indices.tolist()

    for si in top_start:
        for ei in top_end:
            if ei < si:
                continue
            if (ei - si + 1) > max_answer_len:
                continue
            if sequence_ids[si] != 1 or sequence_ids[ei] != 1:
                continue

            sc = int(offset_mapping[si][0])
            ec = int(offset_mapping[ei][1])

            if ec <= sc:
                continue

            score = float(start_logits[si] + end_logits[ei])

            if score > best_score:
                best_score = score
                best_span = (si, ei, sc, ec)

    if best_span is None:
        return "", float("-inf"), 0.0

    si, ei, sc, ec = best_span
    answer = context[sc:ec].strip()

    if not answer:
        return "", float("-inf"), 0.0

    # --- Confidence via sigmoid of the combined logit score ---
    # The null-answer score is typically start_logits[0] + end_logits[0] (CLS).
    null_score = float(start_logits[0] + end_logits[0])
    score_diff = best_score - null_score

    # Sigmoid with a temperature to spread values across [0, 1].
    confidence = float(torch.sigmoid(torch.tensor(score_diff / 3.0)).item())

    # Boost: also factor in the softmax concentration on the answer tokens.
    # With smaller chunks (~200 tokens), softmax values are much more peaked.
    start_prob = float(torch.softmax(start_logits, dim=0)[si])
    end_prob = float(torch.softmax(end_logits, dim=0)[ei])
    geom_mean = float(np.sqrt(start_prob * end_prob))

    # Blend: 70% sigmoid-based, 30% geometric-mean-based.
    confidence = 0.7 * confidence + 0.3 * min(geom_mean * 5.0, 1.0)
    confidence = max(0.0, min(confidence, 1.0))

    return answer, best_score, confidence


def answer_question(question, embed_model, qa_assets, top_k=5):
    """
    Run extractive QA **per-chunk** instead of concatenating chunks.
    This dramatically improves answer quality and confidence scores because
    each chunk is short enough to avoid truncation and softmax dilution.
    """

    if len(embeddings_store) == 0:
        raise ValueError("No paper has been ingested for question answering.")

    question = question.strip()
    if not question:
        raise ValueError("Question cannot be empty.")

    query_vec = embed_model.encode([question])
    sims = cosine_similarity(query_vec, embeddings_store)[0]

    top_k = max(1, min(top_k, len(sims)))
    top_k_idx = np.argsort(sims)[-top_k:]

    qa_tokenizer, qa_model_inst = qa_assets

    # Run QA on each retrieved chunk independently.
    best_answer = ""
    best_confidence = 0.0
    best_overall_score = float("-inf")

    for chunk_idx in reversed(top_k_idx):
        chunk_text = metadata_store[str(chunk_idx)]["text"]

        answer, score, confidence = _qa_single_chunk(
            question, chunk_text, qa_tokenizer, qa_model_inst
        )

        if score > best_overall_score and answer:
            best_overall_score = score
            best_answer = answer
            best_confidence = confidence

    has_answer = bool(best_answer)

    return {
        "answer": best_answer if has_answer else "",
        "confidence": best_confidence,
        "context_chars": sum(
            len(metadata_store[str(i)]["text"]) for i in top_k_idx
        ),
        "has_answer": has_answer,
    }